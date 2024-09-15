import html, json
import pandas as pd
import io
import jwt
import string
import re
from django.conf import settings
from django.http import HttpResponse, FileResponse as FileRes
from rest_framework.views import APIView
from django.utils import timezone
from rest_framework.response import Response
from VectorDB.VectorDB import vector_db
from Query.run_query import answer_docs
from Application.Utilities.RAGCommon import create_collection_name
from Application.serializer.Chat import DocChatSerializers
from Constants.constants import (
    flatten_serializer_errors,
    allowed_citations_headers,
    TypeOfSearch
)
from Application.models import UserAPICall
from Query.run_query import answer_docs
import logging
from django.core.cache import cache
import hashlib
import time
from docx import Document
from docx.shared import Inches
import io
from django.http import FileResponse

logger = logging.getLogger(__name__)
class DocChatClass(APIView):
    def post(self, request):
        start_time = time.time()
        data = request.data
        user_id = request.headers.get("User-ID")
        
        if not user_id:
            return Response({"Message": "User-ID header is required"}, status=400)
        # Check API call frequency
        user_calls, created = UserAPICall.objects.get_or_create(user_id=user_id)
        if not created:
            user_calls.call_count += 1
            user_calls.last_call = timezone.now()
        else:
            user_calls.call_count = 1
        user_calls.save()
        
        if user_calls.call_count > 5:
            return Response({"Message": "Too many requests"}, status=429)
        
        data_main = {
            # "document_id": data.pop("documentID", None),
            "query": data.pop("query", None),
            "type_of_search": data.pop("type_of_search", TypeOfSearch.default),
            "top_k": int(data.get("top_k", 10)),
            "threshold": float(data.get("threshold", 0.5)),
        }
        header = request.headers.get("Org")
        serializer = DocChatSerializers(data=data_main)
        if serializer.is_valid():
            # document_id = serializer.validated_data["document_id"]
            query = serializer.validated_data["query"]
            type_of_search = serializer.validated_data["type_of_search"]
            top_k = serializer.validated_data["top_k"]
            threshold = serializer.validated_data["threshold"]
            collection_name = create_collection_name(header, type_of_search, user_id)
            # document_exists = vector_db.check_file_existence(
            #     document_id, collection_name, type_of_search
            # )
            pattern = r'heading - (.*?) accurately'
            matches = re.search(pattern, query, re.IGNORECASE)
            if matches:
                headings = matches.group(1).split(' and ')
                for heading in headings:
                    heading = heading.strip()
                    # Check if the article exists in the vector database
                    if vector_db.check_article_exists(heading, None, None):
                        # Generate cache key
                        cache_key = self.generate_cache_key(heading, top_k, threshold)
                        
                        # Try to get the results from cache
                        cached_results = cache.get(cache_key)
                        if cached_results:
                            logger.info(f"Cache hit for article: {heading}")
                            return self.create_word_document_response(cached_results)
                        
                        # If not in cache, retrieve the article data
                        article_data = vector_db.get_article_data(heading)
                        if article_data:
                            answer_html_decoded = html.unescape(article_data['content'])
                            answer_plain_text = answer_html_decoded.replace(',', '\n')
                            
                            # Process the content
                            content = self.process_content(answer_plain_text)
                            
                            # Cache the results
                            cache.set(cache_key, content, timeout=3600)  # Cache for 1 hour
                            
                            data = self.create_word_document_response(content)
                            end_time = time.time()
                            inference_time = end_time - start_time
                            logger.info(f"User {user_id} made a chat request. Inference time: {inference_time:.2f}s")
                            return Response(data, status=200)
                        else:
                            return Response({"Message": "Article data not found"}, status=404)
                    else:
                        return Response({"Message": f"Article with heading '{heading}' not found"}, status=404)
            
            # If no matching article found, return an error response
            return Response({"Message": "No matching article found for the given query"}, status=404)
        else:
            return Response(
                status=400,
                data={"Message": flatten_serializer_errors(serializer.errors)},
            )
                    
            # if document_exists:
            #     # Create JWT token
            #     # token_data = {
            #     #     "document_id": document_id,
            #     #     # "chat_id": result_content[0].get("chat_id"),  # Assuming chat_id is part of result_content
            #     #     "exp": timezone.now() + timezone.timedelta(hours=1)  # Token expires in 1 hour
            #     # }
            #     # token = jwt.encode(token_data, settings.SECRET_KEY, algorithm="HS256")
            #     # data['token'] = token  # Include token in response
            #     # # return Response(data, status=200)
                
            #     ###  
                
                
            #     # Generate a cache key based on the search parameters
            #     cache_key = self.generate_cache_key(query, document_id, top_k, threshold)
                
            #     # Try to get the results from cache
            #     cached_results = cache.get(cache_key)
            #     if cached_results:
            #         logger.info(f"Cache hit for query: {query}")
            #         return Response(cached_results)
                
            #     # If not in cache, perform the search
            #     result, result_content, citation = answer_docs(
            #         query, document_id, collection_name, header, type_of_search, top_k, threshold
            #     )
            #     if header in allowed_citations_headers:
            #         data = {
            #             "query": query,
            #             "answer": result,
            #             "source": result_content,
            #             "citation": citation,
            #             "time": timezone.localtime(timezone.now()),
            #         }
                    
                    
            #     else:
            #         data = {
            #             "query": query,
            #             "answer": result,
            #             "source": result_content[0]["text"],
            #             "time": timezone.localtime(timezone.now()),
            #         }
                    
                    
            #         answer_html_decoded = html.unescape(data["answer"])
            #         # print(answer_html_decoded)
            #         # Extract headings from the query
            #         pattern = r'heading - (.*?) accurately'
            #         matches = re.search(pattern, data["query"], re.IGNORECASE)
            #         if matches:
            #             # Split the matched group by 'and' to handle multiple headings
            #             headings = matches.group(1).split(' and ')
            #             # Remove each heading and its following colon from the answer_html_decoded
            #             for heading in headings:
            #                 # Regex to match the heading and remove it along with any trailing data structure syntax
            #                 answer_html_decoded = re.sub(r'"' + re.escape(heading.strip()) + r'"\s*:\s*', '', answer_html_decoded)
                    
                    
            #         ###testing testing vvvv 
            #         print(answer_html_decoded)
                    
                    
            #         answer_plain_text = answer_html_decoded
            #         answer_plain_text = answer_plain_text.replace(',', '\n')
                    
            #         # Split the text into lines
            #         lines = answer_plain_text.split('\n')
                    
            #         # Group the content under headings
            #         content = {}
            #         current_heading = None
            #         for line in lines:
            #             if ':' in line:
            #                 key, value = line.split(':', 1)
            #                 key = key.strip()
            #                 value = value.strip()
            #                 if value:  # If there's a value, it's a subheading
            #                     if current_heading:
            #                         content[current_heading].append(f"{key}: {value}")
            #                     else:
            #                         content[key] = [value]
            #                 else:  # If there's no value, it's a main heading
            #                     current_heading = key
            #                     content[current_heading] = []
            #             elif current_heading:
            #                 content[current_heading].append(line.strip())
                    
            #         # Create a Word document
            #         doc = Document()
                    
            #         # Add content to the document
            #         for heading, items in content.items():
            #             doc.add_heading(heading, level=1)
            #             for item in items:
            #                 doc.add_paragraph(item)
                    
            #         # Save the document to a BytesIO object
            #         buffer = io.BytesIO()
            #         doc.save(buffer)
            #         buffer.seek(0)
                    
            #         # Create the response with the Word document
            #         response = FileResponse(buffer, as_attachment=True, filename='chatbot_result.docx')
            #         return response
            #         # # Ensure the string is properly formatted as JSON
            #         # # Remove all curly braces first to avoid nested or mismatched braces
            #         # answer_html_decoded = re.sub(r'[{}]', '', answer_html_decoded)
            #         # # Now enclose the entire string in a single pair of curly braces
            #         # answer_html_decoded = '{' + answer_html_decoded + '}'
            #         # if not (answer_html_decoded.startswith('{') and answer_html_decoded.endswith('}')):
            #         #     answer_html_decoded = '{' + answer_html_decoded + '}'
            #         # answer_html_decoded = answer_html_decoded.replace("<br>", '')
            #         # #to remove the brackets from the first and last of the string and then remove the first word from the string and make it as the heading of the excel file and remove it's quotes
            #         # # answer_html_decoded = answer_html_decoded[1:-1]
            #         # #remove the first word encountered and make it as the heading of the excel table
            #         # # Remove the first word and the colon following it
            #         # # first_colon_index = answer_html_decoded.find(':')
            #         # # if first_colon_index != -1:
            #         # #     answer_html_decoded = answer_html_decoded[first_colon_index + 1:].strip()
            #         # print(answer_html_decoded)
            #         # answer_dict = json.loads(answer_html_decoded)
                    
            #         # # df = pd.DataFrame([answer_dict]).transpose() # Convert dictionary to DataFrame
            #         # df = pd.DataFrame(list(answer_dict.items()), columns=['Sub-heading', 'Value'])
                    
            #         # buffer = io.BytesIO()
            #         # with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            #         #     df.to_excel(writer, index = False)
            #         # buffer.seek(0)
            #         # response = FileRes(buffer, as_attachment=True, filename='chatbot_result.xlsx')
            #         # return response
                    
            #     # Cache the results
            #     cache.set(cache_key, data, timeout=3600)  # Cache for 1 hour
                
            #     end_time = time.time()
            #     inference_time = end_time - start_time
            #     logger.info(f"User {user_id} made a chat request. Inference time: {inference_time:.2f}s")
                
            #     return Response(data, status=200)
            # else:
            #     return Response(
            #         {"Message": "The news articles are not available at this moment"},
            #         status=400,
            #     )
            # else:
            # return Response(
            #     status=400,
            #     data={"Message": flatten_serializer_errors(serializer.errors)},
            # )
    
    def generate_cache_key(self, query, top_k, threshold):
        # Create a unique cache key based on the search parameters
        key_parts = f"{query}:{top_k}:{threshold}"
        return f"chat_results:{hashlib.md5(key_parts.encode()).hexdigest()}"
    
    def process_content(self, text):
        lines = text.split('\n')
        content = {}
        current_heading = None
        for line in lines:
            if ':' in line:
                key, value = line.split(':', 1)
                key = key.strip()
                value = value.strip()
                if value:  # If there's a value, it's a subheading
                    if current_heading:
                        content[current_heading].append(f"{key}: {value}")
                    else:
                        content[key] = [value]
                else:  # If there's no value, it's a main heading
                    current_heading = key
                    content[current_heading] = []
            elif current_heading:
                content[current_heading].append(line.strip())
        return content

    def create_word_document_response(self, content):
        doc = Document()
        for heading, items in content.items():
            doc.add_heading(heading, level=1)
            for item in items:
                doc.add_paragraph(item)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return FileResponse(buffer, as_attachment=True, filename='chatbot_result.docx')
#writing an get api to convert the result (only the from "answer" section) obtained from the chatbot to an excel file and return the file to the user
class ResultFileClass(APIView):
    def get(self, request):
        token = request.headers.get('Authorization', '').split(' ')[1]  # Get token from header
        try:
            token_data = jwt.decode(token, settings.SECRET_KEY, algorithms=["HS256"])
        except jwt.ExpiredSignatureError:
            return Response({"Message": "Token expired"}, status=403)
        except jwt.InvalidTokenError:
            return Response({"Message": "Invalid token"}, status=403)

        document_id = token_data['document_id']
        chat_id = token_data['resultID']
        data_main = {
            "document_id": request.query_params.get("documentID", None),
            "chat_id": request.query_params.get("resultID", None),
            "type_of_search": request.query_params.get("type_of_search", TypeOfSearch.default),
        }
        header = request.headers.get("Org")
        serializer = DocChatSerializers(data=data_main)
        if serializer.is_valid():
            document_id = serializer.validated_data["document_id"]
            chat_id = serializer.validated_data["resultID"]
            type_of_search = serializer.validated_data["type_of_search"]
            collection_name = create_collection_name(header, type_of_search)
            document_exists = vector_db.check_file_existence(
                document_id, collection_name, type_of_search
            )
            if document_exists:
                chat_data = vector_db.get_chat_data(chat_id)
                if chat_data:
                    chat_data = chat_data[0]
                    # Decode HTML entities and parse JSON
                    answer_html_decoded = html.unescape(chat_data["answer"])
                    answer_dict = json.loads(answer_html_decoded)
                    df = pd.DataFrame([answer_dict])  # Convert dictionary to DataFrame
                    output = io.BytesIO()
                    writer = pd.ExcelWriter(output, engine='xlsxwriter')
                    df.to_excel(writer, sheet_name='Sheet1')
                    writer.save()
                    output.seek(0)
                    response = HttpResponse(output, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                    response['Content-Disposition'] = 'attachment; filename=chatbot_result.xlsx'
                    return response
                else:
                    return Response(
                        {"Message": "Chat data is not available"},
                        status=400,
                    )
            else:
                return Response(
                    {"Message": "Document is not available"},
                    status=400,
                )
        else:
            return Response(
                status=400,
                data={"Message": flatten_serializer_errors(serializer.errors)},
            )
    
from pymilvus import connections

try:
    connections.connect("default", host='127.0.0.1', port='19530')
    print("Connected to Milvus server.")
except Exception as e:
    print(f"Failed to connect to Milvus server: {e}")